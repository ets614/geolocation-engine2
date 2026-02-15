#!/usr/bin/env python3
"""Export PROGRESS.md to professional DOCX format for business audience."""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def read_progress_markdown(file_path: str) -> str:
    """Read PROGRESS.md file."""
    with open(file_path, 'r') as f:
        return f.read()


def create_progress_docx(markdown_path: str, output_path: str = "docs/feature/ai-detection-cop-integration/PROGRESS.docx"):
    """Convert PROGRESS.md to professional DOCX document."""

    content = read_progress_markdown(markdown_path)
    lines = content.split('\n')

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ====== TITLE ======
    title = doc.add_heading('AI Detection to CoP Integration', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    title_run.font.size = Pt(24)

    subtitle = doc.add_paragraph('Project Progress Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()  # Spacing

    # ====== CONTENT PROCESSING ======
    in_code_block = False
    skip_next = False

    for i, line in enumerate(lines):
        if skip_next:
            skip_next = False
            continue

        # Skip empty lines at start
        if not doc.paragraphs or (doc.paragraphs[-1].text == '' and i < 10):
            continue

        # Handle headings
        if line.startswith('# '):
            title_text = line.replace('# 🎯 ', '').replace('# ', '').strip()
            if title_text:
                heading = doc.add_heading(title_text, level=1)
                heading_run = heading.runs[0]
                heading_run.font.color.rgb = RGBColor(0, 102, 204)
            continue

        if line.startswith('## '):
            heading_text = line.replace('## ', '').replace('📊 ', '').replace('🏗 ', '').replace('🎁 ', '').replace('📈 ', '').replace('✨ ', '').replace('🚀 ', '').strip()
            if heading_text:
                heading = doc.add_heading(heading_text, level=2)
                heading_run = heading.runs[0]
                heading_run.font.color.rgb = RGBColor(0, 102, 204)
            continue

        # Handle code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            if not in_code_block:
                doc.add_paragraph()  # Add spacing after code block
            continue

        if in_code_block:
            # Add as monospace code
            p = doc.add_paragraph(line or '', style='Normal')
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(64, 64, 64)
            p.paragraph_format.left_indent = Inches(0.25)
            continue

        # Handle regular lines
        if line.strip():
            # Skip table separator lines
            if line.startswith('|') or line.startswith('─'):
                continue

            # Clean up text
            text = line
            for emoji in ['🎯', '📊', '🏗', '🎁', '📈', '✨', '🚀', '✅', '⏳', '⏭', '📄', '📱', '📋', '⏰', '🧪', '🎨', '🔄']:
                text = text.replace(emoji, '')

            text = text.strip()

            if text and not text.startswith('='):
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(6)

    # ====== FOOTER ======
    doc.add_paragraph()  # Spacing
    footer_para = doc.add_paragraph()
    footer_para.paragraph_format.border_pt = 1
    footer_run = footer_para.add_run('_' * 100)
    footer_run.font.size = Pt(8)

    # Footer text
    gen_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    footer_text = doc.add_paragraph()
    footer_text_run = footer_text.add_run(f"Generated: {gen_time}")
    footer_text_run.font.size = Pt(9)
    footer_text_run.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    doc.save(output_path)

    size = os.path.getsize(output_path)
    return {
        'path': output_path,
        'size': size,
        'generated': datetime.now().isoformat()
    }


if __name__ == '__main__':
    import sys

    markdown_file = sys.argv[1] if len(sys.argv) > 1 else \
        'docs/feature/ai-detection-cop-integration/PROGRESS.md'

    output_file = sys.argv[2] if len(sys.argv) > 2 else \
        'docs/feature/ai-detection-cop-integration/PROGRESS.docx'

    try:
        result = create_progress_docx(markdown_file, output_file)
        print(f"✅ DOCX exported successfully")
        print(f"📄 File: {result['path']}")
        print(f"📊 Size: {result['size']:,} bytes")
        print(f"⏰ Generated: {result['generated']}")
    except FileNotFoundError:
        print(f"❌ Error: Could not find {markdown_file}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error exporting DOCX: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
